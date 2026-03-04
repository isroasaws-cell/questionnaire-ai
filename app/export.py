from docx import Document


def export_docx(answers):

    document = Document()

    document.add_heading("Questionnaire Answers", level=1)

    for answer in answers:

        document.add_heading("Question", level=2)
        document.add_paragraph(answer.question_text)

        document.add_heading("Answer", level=3)
        document.add_paragraph(answer.answer_text)

        document.add_paragraph("Citations:")
        document.add_paragraph(answer.citations)

        document.add_paragraph(f"Confidence: {answer.confidence_score}")

        document.add_paragraph("----------------------------")

    file_path = "exports/output.docx"

    document.save(file_path)

    return file_path